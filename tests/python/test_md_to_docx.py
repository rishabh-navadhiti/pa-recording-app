"""Golden-structure tests for md_to_docx.py (Phase 5).

Asserts the *structure* of the produced .docx (paragraph styles, run bold/
underline, table rows/cells/shading) rather than byte-comparing the binary —
robust to zip ordering/timestamps. Per decision A7 these are the precondition
for ever porting md_to_docx to Node: they pin the load-bearing formatting so a
port can be proven equivalent.

Run: python -m unittest discover -s tests/python
"""

import os
import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..', 'python'))

try:
    import md_to_docx
    from docx import Document
    from docx.oxml.ns import qn
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False


MARKDOWN = "\n".join([
    "# Title One",
    "## Section Two",
    "",
    "- bullet alpha",
    "  - sub bullet",
    "1. numbered one",
    "",
    "Normal text with **bold word** and <u>underlined</u> and **<u>both</u>**.",
    "",
    "WORK RESTRICTIONS:",
    "DOB: 12/26/1993",
    "",
    "| Code | Description |",
    "| --- | --- |",
    "| M25.511 | Pain in right shoulder |",
    "",
    "<!-- layout -->",
    "| TO: | Dr. Smith |",
    "| --- | --- |",
    "| Patient: | Jane Doe |",
    "",
])


@unittest.skipUnless(HAVE_DOCX, "python-docx not installed")
class MdToDocxGoldenTest(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        cls.tmp = tempfile.mkdtemp()
        md_path = Path(cls.tmp) / "note.md"
        md_path.write_text(MARKDOWN, encoding="utf-8")
        cls.docx_path = Path(cls.tmp) / "note.docx"
        md_to_docx.convert_md_to_docx(md_path, cls.docx_path)
        cls.doc = Document(str(cls.docx_path))

    def _para(self, contains):
        for p in self.doc.paragraphs:
            if contains in p.text:
                return p
        self.fail(f"no paragraph containing {contains!r}")

    @staticmethod
    def _has_shading(cell):
        tcPr = cell._tc.find(qn('w:tcPr'))
        return tcPr is not None and tcPr.find(qn('w:shd')) is not None

    def test_heading_styles(self):
        self.assertEqual(self._para("Title One").style.name, "Heading 1")
        self.assertEqual(self._para("Section Two").style.name, "Heading 2")

    def test_list_styles(self):
        self.assertEqual(self._para("bullet alpha").style.name, "List Bullet")
        self.assertEqual(self._para("sub bullet").style.name, "List Bullet 2")
        self.assertEqual(self._para("numbered one").style.name, "List Number")

    def test_inline_bold_and_underline(self):
        runs = {r.text: r for r in self._para("Normal text with").runs}
        self.assertTrue(runs["bold word"].bold)
        self.assertTrue(runs["underlined"].underline)
        self.assertTrue(runs["both"].bold)
        self.assertTrue(runs["both"].underline)

    def test_allcaps_heading_is_bolded(self):
        self.assertTrue(any(r.bold for r in self._para("WORK RESTRICTIONS:").runs))

    def test_field_line_with_digits_is_not_force_bolded(self):
        # "DOB: 12/26/1993" must NOT be treated as an ALL-CAPS heading (it has
        # digits) — so it stays a plain, unbolded paragraph.
        self.assertFalse(any(r.bold for r in self._para("DOB: 12/26/1993").runs))

    def test_gfm_table_structure_and_shaded_header(self):
        t = self.doc.tables[0]
        self.assertEqual((len(t.rows), len(t.columns)), (2, 2))
        self.assertEqual(t.cell(0, 0).text, "Code")
        self.assertEqual(t.cell(1, 0).text, "M25.511")
        self.assertTrue(self._has_shading(t.cell(0, 0)), "GFM header cell is shaded")

    def test_layout_table_is_unshaded(self):
        t = self.doc.tables[1]
        self.assertEqual(len(t.columns), 2)
        self.assertEqual(t.cell(0, 0).text, "TO:")
        self.assertEqual(t.cell(1, 1).text, "Jane Doe")
        self.assertFalse(self._has_shading(t.cell(0, 0)), "layout table is borderless/unshaded")


if __name__ == '__main__':
    unittest.main()
