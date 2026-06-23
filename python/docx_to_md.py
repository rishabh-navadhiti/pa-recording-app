#!/usr/bin/env python3
"""
docx_to_md.py — Convert a Word (.docx) SOAP / clinical note to Markdown.

The faithful inverse of the recording-app's python/md_to_docx.py. Where
md_to_docx.py renders markdown → styled Word, this reads the Word styling
(heading styles, dark-blue headings, bold, underline, tables, colour) back
into the markdown vocabulary that the app's pipeline and the cdi-costigan
skill expect:

    Word heading style (Heading 1..4)           → #, ##, ###, ####
    Dark-blue bold heading run (e.g. 1F497D),
      not a Word heading style                  → ## (treated as a section heading)
    Bold + underline section label              → bold+underline preserved as **<u>..</u>**
                                                   (and, if it's a standalone label line,
                                                    promoted to ### so it reads as a sub-heading)
    Grey/secondary sub-label run                 → kept as bold text
    ALL CAPS line ending in ':'                  → left as a plain ALL-CAPS line
                                                   (md_to_docx re-bolds these)
    **bold** inline                              → **bold**
    underline inline                             → <u>..</u>
    Word table                                   → GFM pipe table (| col | col |)
    list paragraph (bullet/number)               → - item  /  1. item  (with indent)
    blank paragraph                              → blank line

Usage:
    python3 docx_to_md.py <path_to_note.docx> [<output.md>]

Output:
    If <output.md> is omitted, writes <stem>.md alongside the input.
    Prints the absolute path of the saved .md to stdout on success.
    Exits 1 on error.

Notes / limits:
    - Reads paragraphs and tables in document order (body-level).
    - Run-level bold/underline are detected from explicit run formatting AND
      from the run's style (so style-driven bold is not lost).
    - "Dark-blue heading" detection keys on the colour used by md_to_docx
      (1F497D) and close variants, since that's how the source notes render
      HISTORY OF PRESENT ILLNESS / PHYSICAL EXAM / ASSESSMENT & PLAN.
    - Inline images, footnotes, headers/footers, and text boxes are ignored.
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    import subprocess
    subprocess.check_call(
        [sys.executable, "-m", "pip", "install", "python-docx", "--break-system-packages", "-q"],
        stdout=subprocess.DEVNULL,
    )
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph


# ---------------------------------------------------------------------------
# Colour / heading detection
# ---------------------------------------------------------------------------

# The dark blue md_to_docx.py uses for headings, plus the navy the source
# Costigan notes use for HISTORY OF PRESENT ILLNESS / PHYSICAL EXAM etc.
HEADING_BLUES = {"1F497D", "000080", "1F3864", "002060", "0000FF"}
# A grey used for secondary/sub labels (e.g. the faint "Past Surgical History"
# and "Impression" sub-labels). Treated as de-emphasised — not a heading.
GREY_PREFIXES = ("80", "808080", "7F7F7F", "A6A6A6", "BFBFBF", "595959", "767171")


def _hex_upper(color):
    """Return the run colour as an UPPER hex string, or None."""
    try:
        rgb = color.rgb
    except Exception:
        return None
    if rgb is None:
        return None
    return str(rgb).upper()


def run_is_bold(run):
    """Bold if the run says so, or its style/parent style does."""
    if run.bold:
        return True
    # style-driven bold
    try:
        st = run.style
        while st is not None:
            f = getattr(st, "font", None)
            if f is not None and f.bold:
                return True
            st = getattr(st, "base_style", None)
    except Exception:
        pass
    return False


def run_is_underline(run):
    if run.underline:
        return True
    try:
        st = run.style
        while st is not None:
            f = getattr(st, "font", None)
            if f is not None and f.underline:
                return True
            st = getattr(st, "base_style", None)
    except Exception:
        pass
    return False


def run_color_hex(run):
    """Explicit run colour hex (upper), or None."""
    try:
        return _hex_upper(run.font.color)
    except Exception:
        return None


def run_text(run):
    """Run text with tabs normalised to a single space, preserving the string."""
    t = run.text or ""
    return t.replace("\t", " ")


# ---------------------------------------------------------------------------
# Inline markdown assembly
# ---------------------------------------------------------------------------

def _wrap(text, bold, underline):
    """Wrap a contiguous span of identically-formatted text in md markers."""
    if not text:
        return ""
    # Preserve leading/trailing whitespace OUTSIDE the markers so we don't
    # produce '** bold **' (which markdown won't render as bold).
    lead = text[: len(text) - len(text.lstrip())]
    trail = text[len(text.rstrip()):]
    core = text.strip()
    if not core:
        return text  # whitespace-only span
    if underline:
        core = f"<u>{core}</u>"
    if bold:
        core = f"**{core}**"
    return f"{lead}{core}{trail}"


def paragraph_inline_md(para):
    """
    Build the inline-markdown string for a paragraph by walking its runs and
    coalescing adjacent runs that share bold/underline state. Returns the
    markdown text (no leading bullet/heading markers).
    """
    spans = []  # list of (text, bold, underline)
    for run in para.runs:
        txt = run_text(run)
        if txt == "":
            continue
        b = run_is_bold(run)
        u = run_is_underline(run)
        if spans and spans[-1][1] == b and spans[-1][2] == u:
            spans[-1][0] += txt
        else:
            spans.append([txt, b, u])

    # Coalesce: if EVERY non-space span is bold, we still emit per-span so a
    # mixed line like "**VAS:** 8/10" renders correctly.
    out = []
    for txt, b, u in spans:
        out.append(_wrap(txt, b, u))
    md = "".join(out)
    # Collapse a fully-bolded line's redundant adjacent markers: '**a****b**' → '**ab**'
    md = md.replace("****", "")
    # Tidy multiple spaces introduced by tab normalisation.
    md = re.sub(r"[ \t]{2,}", " ", md)
    return md.rstrip()


# ---------------------------------------------------------------------------
# Paragraph classification
# ---------------------------------------------------------------------------

def heading_level_from_style(para):
    """Return 1..4 if the paragraph uses a Word Heading style, else None."""
    name = (para.style.name or "") if para.style else ""
    m = re.match(r"Heading (\d)", name)
    if m:
        lvl = int(m.group(1))
        return min(max(lvl, 1), 4)
    if name == "Title":
        return 1
    return None


def list_info(para):
    """
    Return ('bullet'|'number', indent_level) if the paragraph is a list item,
    else (None, 0). Detects via numbering properties (numPr) and style name.
    """
    style_name = (para.style.name or "") if para.style else ""
    is_list_style = "List" in style_name or style_name.startswith("List")
    numPr = para._p.find(qn("w:pPr") + "/" + qn("w:numPr")) if para._p is not None else None
    # python-docx: numPr lives under pPr; search robustly
    has_numbering = False
    ilvl = 0
    pPr = para._p.find(qn("w:pPr"))
    if pPr is not None:
        np = pPr.find(qn("w:numPr"))
        if np is not None:
            has_numbering = True
            il = np.find(qn("w:ilvl"))
            if il is not None:
                try:
                    ilvl = int(il.get(qn("w:val")))
                except Exception:
                    ilvl = 0
    if not (has_numbering or is_list_style):
        return (None, 0)
    kind = "number" if ("Number" in style_name) else "bullet"
    return (kind, ilvl)


def para_indent_spaces(para):
    """Approximate leading-space indent from the paragraph's left indent (twips→spaces)."""
    try:
        ind = para.paragraph_format.left_indent
        if ind is not None:
            inches = ind.inches
            # ~0.25" per indent step → 2 spaces per step (keeps sub-bullets detectable)
            steps = int(round(inches / 0.25))
            return max(0, steps) * 2
    except Exception:
        pass
    return 0


def plain_text(para):
    return "".join(run_text(r) for r in para.runs).strip()


def is_allcaps_heading(text):
    """Same shape md_to_docx.is_allcaps_heading recognises (so we leave it plain)."""
    stripped = text.strip()
    if len(stripped) < 4:
        return False
    return bool(re.match(r"^[A-Z][A-Z\s/&,().:-]+:?\s*$", stripped))


def dominant_color_and_bold(para):
    """Return (color_hex_or_None, all_runs_bold) over the non-empty runs."""
    colors = set()
    bolds = []
    any_text = False
    for run in para.runs:
        if run_text(run).strip() == "":
            continue
        any_text = True
        c = run_color_hex(run)
        if c:
            colors.add(c)
        bolds.append(run_is_bold(run))
    if not any_text:
        return (None, False)
    color = next(iter(colors)) if len(colors) == 1 else None
    return (color, all(bolds) and len(bolds) > 0)


# ---------------------------------------------------------------------------
# Table rendering → GFM
# ---------------------------------------------------------------------------

def _iter_cell_blocks(cell):
    """Yield ('p', Paragraph) / ('t', Table) for a cell's children in order.

    python-docx's cell.paragraphs ignores nested tables, so we walk the cell's
    XML children directly. Nested tables are common in these EMR exports — the
    Past Surgical History / Medications / Diagnosis blocks wrap their real data
    rows in a table nested inside a single outer cell. Without this recursion
    that content is silently dropped.
    """
    for child in cell._tc.iterchildren():
        if child.tag == qn("w:p"):
            yield ("p", Paragraph(child, cell))
        elif child.tag == qn("w:tbl"):
            yield ("t", Table(child, cell))


def cell_has_nested_table(cell):
    return any(kind == "t" for kind, _ in _iter_cell_blocks(cell))


def _cell_md(cell):
    """Inline markdown for a table cell — paragraphs joined; nested-table text flattened in.

    For a plain cell this is just its paragraphs. If the cell contains a nested
    table (layout wrapper), the nested rows are flattened into a single string
    so no content is lost even when the cell is rendered inline.
    """
    parts = []
    for kind, blk in _iter_cell_blocks(cell):
        if kind == "p":
            txt = paragraph_inline_md(blk).strip()
            if txt:
                parts.append(txt)
        else:  # nested table → flatten its rows to "a | b | c" lines
            for r in build_grid(blk):
                r = [c for c in _dedupe_merged(r) if c.strip()]
                if r:
                    parts.append(" — ".join(r))
    return " ".join(parts).replace("|", "\\|").strip()


def _dedupe_merged(cells):
    """
    Collapse consecutive identical cells in a row. Word horizontal merges
    (gridSpan) make python-docx repeat the merged cell's text across each
    spanned column; this restores a single logical cell so a row like
    'Tobacco Use | Tobacco Use | Tobacco Use' becomes just 'Tobacco Use'.
    """
    out = []
    for c in cells:
        if out and out[-1] == c and c != "":
            continue
        out.append(c)
    return out


def build_grid(table):
    """Return the table's cell-text grid (rows of inline-md strings)."""
    grid = []
    for row in table.rows:
        grid.append([_cell_md(cell) for cell in row.cells])
    return grid


def _first_nonempty_para(table):
    """The first paragraph in the table that has visible text (for heading detection)."""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if plain_text(p):
                    return p
    return None


def expand_table_rows(table):
    """
    Return the table's logical rows (lists of inline-md cell strings),
    recursing into nested tables.

    EMR exports wrap real data grids (Past Surgical History, Medications,
    Diagnosis) in a table nested inside a single outer cell. python-docx's
    .rows API does not recurse, so those rows would be lost. Here, when a row
    is effectively just a wrapper around a nested table, we splice the nested
    table's expanded rows in at that point instead of emitting the wrapper.
    """
    out = []
    for row in table.rows:
        # Gather this row's cells as (text, nested_tables[]) so we can decide
        # whether the row is a real data row or a wrapper around a nested table.
        cell_texts = []
        nested_in_row = []
        for cell in row.cells:
            nested = [blk for kind, blk in _iter_cell_blocks(cell) if kind == "t"]
            para_txt = " ".join(
                paragraph_inline_md(p).strip()
                for p in cell.paragraphs
                if paragraph_inline_md(p).strip()
            ).replace("|", "\\|").strip()
            cell_texts.append(para_txt)
            nested_in_row.extend(nested)

        deduped = _dedupe_merged(cell_texts)
        row_has_text = any(c.strip() for c in deduped)

        # If the row carries no own text but wraps nested table(s), splice the
        # nested rows in directly (the wrapper cell is just a layout container).
        if nested_in_row and not row_has_text:
            for nt in nested_in_row:
                out.extend(expand_table_rows(nt))
            continue

        # Normal data row. If it also has nested tables (rare), append them after.
        out.append(deduped)
        for nt in nested_in_row:
            out.extend(expand_table_rows(nt))

    return out


def table_to_md(table):
    """Render a Word table as a GFM pipe table. First row = header. Recurses into nested tables."""
    grid = expand_table_rows(table)
    grid = [r for r in grid if any(c.strip() for c in r)]  # drop blank rows
    if not grid:
        return []

    ncol = max(len(r) for r in grid)
    grid = [r + [""] * (ncol - len(r)) for r in grid]

    rows_md = []
    header = grid[0]
    rows_md.append("| " + " | ".join(header) + " |")
    rows_md.append("| " + " | ".join(["---"] * ncol) + " |")
    for r in grid[1:]:
        rows_md.append("| " + " | ".join(r) + " |")
    return rows_md


# ---------------------------------------------------------------------------
# Body iteration in document order
# ---------------------------------------------------------------------------

def iter_block_items(parent):
    """Yield Paragraph and Table objects in document order."""
    from docx.document import Document as _Doc
    if isinstance(parent, _Doc):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._element
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


# ---------------------------------------------------------------------------
# Main conversion
# ---------------------------------------------------------------------------

def convert_docx_to_md(docx_path: Path, md_path: Path):
    doc = Document(str(docx_path))
    out_lines = []

    def emit(line=""):
        out_lines.append(line)

    prev_blank = True  # treat start as already-blank to avoid leading blank line

    for block in iter_block_items(doc):
        if isinstance(block, Table):
            # Expand the table — recursing into nested tables — then drop blank
            # rows so the true logical shape is visible.
            grid = [r for r in expand_table_rows(block) if any(c.strip() for c in r)]
            logical_cols = max((len(r) for r in grid), default=0)
            flat = [c for r in grid for c in r if c.strip()]

            if not prev_blank:
                emit("")

            # Single-cell table holding a blue/ALL-CAPS heading → ## section heading.
            first_para = _first_nonempty_para(block)
            if len(flat) == 1 and first_para is not None:
                color, all_bold = dominant_color_and_bold(first_para)
                htext = plain_text(first_para)
                if (color in HEADING_BLUES and all_bold) or is_allcaps_heading(htext):
                    emit("## " + htext)
                    emit("")
                    prev_blank = True
                    continue

            # Genuinely single-column block (a label/short note, no real grid) →
            # one line per non-empty cell rather than a degenerate 1-col table.
            if logical_cols <= 1:
                for r in grid:
                    for c in r:
                        if c.strip():
                            emit(c)
                emit("")
                prev_blank = True
                continue

            # Real ≥2-column data table (including nested data tables spliced in
            # by expand_table_rows, e.g. Past Surgical History) → GFM table.
            out_lines.extend(table_to_md(block))
            emit("")
            prev_blank = True
            continue

        para = block
        text = plain_text(para)

        # Blank paragraph → single blank line (collapse runs of blanks)
        if text == "":
            if not prev_blank:
                emit("")
                prev_blank = True
            continue

        # 1) Word heading style → #.. by level
        lvl = heading_level_from_style(para)
        if lvl is not None:
            emit("#" * lvl + " " + text)
            emit("")
            prev_blank = True
            continue

        # 2) List item → - / 1.  with indent
        kind, ilvl = list_info(para)
        if kind is not None:
            indent = "  " * ilvl
            inline = paragraph_inline_md(para) or text
            marker = "1." if kind == "number" else "-"
            emit(f"{indent}{marker} {inline}")
            prev_blank = False
            continue

        color, all_bold = dominant_color_and_bold(para)

        # 3) Dark-blue heading run (not a Word heading style) → ## section heading.
        #    This is how the source notes render HISTORY OF PRESENT ILLNESS / PHYSICAL EXAM.
        if color in HEADING_BLUES and all_bold:
            if not prev_blank:
                emit("")
            # Keep the literal text (often ALL CAPS with a trailing colon).
            emit("## " + text)
            emit("")
            prev_blank = True
            continue

        # 4) Standalone bold+underline label line → ### sub-heading,
        #    preserving the underline marker so a round-trip keeps the style.
        #    (e.g. "Past Surgical History:", "Diagnostic Studies History:", "Radiographs:")
        if all_bold and _line_all_underlined(para) and _looks_like_label(text):
            if not prev_blank:
                emit("")
            # Promote to a sub-heading. The ### already implies bold/blue styling
            # on the md→docx round-trip, so drop the redundant **<u>..</u>** markers
            # and emit the plain label text.
            emit("### " + text)
            emit("")
            prev_blank = True
            continue

        # 5) ALL CAPS heading-ish line → leave plain (md_to_docx re-bolds it).
        if is_allcaps_heading(text) and not all_bold:
            if not prev_blank:
                emit("")
            emit(text)
            prev_blank = False
            continue

        # 6) Normal paragraph (with inline **bold** / <u>underline</u>).
        inline = paragraph_inline_md(para)
        # If the whole line is bold (a sub-label like a grey "Impression"),
        # inline already carries the ** markers; keep as-is.
        emit(inline if inline else text)
        prev_blank = False

    # Trim trailing blank lines, ensure single trailing newline.
    while out_lines and out_lines[-1] == "":
        out_lines.pop()
    md = "\n".join(out_lines) + "\n"
    md_path.write_text(md, encoding="utf-8")


def _line_all_underlined(para):
    """True if every non-space run in the paragraph is underlined."""
    saw = False
    for run in para.runs:
        if run_text(run).strip() == "":
            continue
        saw = True
        if not run_is_underline(run):
            return False
    return saw


def _looks_like_label(text):
    """A short heading-like label, typically ending in ':' (e.g. 'Past Surgical History:')."""
    t = text.strip()
    if len(t) > 60:
        return False
    # Title-case-ish or ends with a colon; not a full sentence.
    return t.endswith(":") or (t == t.title()) or t.isupper()


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) not in (2, 3):
        print("Usage: python3 docx_to_md.py <path_to_note.docx> [<output.md>]", file=sys.stderr)
        sys.exit(1)

    docx_path = Path(sys.argv[1]).absolute()
    if not docx_path.exists():
        print(f"ERROR: File not found: {docx_path}", file=sys.stderr)
        sys.exit(1)
    if docx_path.suffix.lower() != ".docx":
        print(f"ERROR: Expected a .docx file, got: {docx_path.name}", file=sys.stderr)
        sys.exit(1)

    md_path = Path(sys.argv[2]).absolute() if len(sys.argv) == 3 else docx_path.with_suffix(".md")

    try:
        convert_docx_to_md(docx_path, md_path)
    except Exception as e:
        print(f"ERROR: Conversion failed: {e}", file=sys.stderr)
        sys.exit(1)

    print(str(md_path))


if __name__ == "__main__":
    main()
