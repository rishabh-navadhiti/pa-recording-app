#!/usr/bin/env python3
"""
md_to_docx.py — Convert a SOAP note Markdown file to a formatted Word (.docx) document.

Usage:
    python3 md_to_docx.py <path_to_soap_note.md>

Output:
    Saves <stem>.docx alongside the input file.
    Prints the absolute path of the saved .docx to stdout on success.
    Exits with code 1 on error and prints the error message.

Supported markdown elements:
    # H1, ## H2, ### H3, #### H4   — Word heading styles (Calibri, dark blue)
    ---                             — Thin horizontal rule
    - item                          — Bullet list (List Bullet)
      - sub-item                    — Indented bullet (List Bullet 2)
    1. item                         — Numbered list (List Number)
    **bold**                        — Inline bold (within any paragraph type)
    ALL CAPS HEADING:               — Auto-bolded (e.g. WC section headings)
    | col | col |                   — GFM pipe table with `|---|---|` separator row
    blank line                      — Paragraph spacer (small vertical gap)
    anything else                   — Normal paragraph
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    subprocess.check_call(
        [sys.executable, "-m", "pip", "install", "python-docx", "--break-system-packages", "-q"],
        stdout=subprocess.DEVNULL
    )
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_horizontal_rule(doc):
    """Add a thin horizontal rule using a bottom border on an empty paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def parse_inline_bold(paragraph, text):
    """
    Add runs to a paragraph, handling **bold** spans.
    Supports: **Key:** value  or  plain text with **bold** words.
    """
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            paragraph.add_run(part)


def is_table_row(line):
    """Pipe-delimited row: starts and ends with '|' and has at least one inner '|'."""
    stripped = line.strip()
    return stripped.startswith('|') and stripped.endswith('|') and stripped.count('|') >= 2


def is_table_separator(line):
    """GFM separator row like `|---|:---:|---:|`. Cells are dashes with optional `:` alignment markers."""
    stripped = line.strip()
    if not is_table_row(stripped):
        return False
    cells = [c.strip() for c in stripped.strip('|').split('|')]
    if not cells:
        return False
    return all(re.match(r'^:?-{2,}:?$', c) for c in cells)


def split_table_row(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]


def set_cell_shading(cell, hex_fill):
    """Apply a solid background colour to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def set_cell_borders(cell, color='BFBFBF', size='4'):
    """Add thin borders on all four sides of a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_md_table(doc, header_row, body_rows):
    """
    Render a GFM markdown table as a styled Word table:
      - dark-blue header bar with bold white text
      - light-blue body rows with thin grey borders
      - inline **bold** is honoured inside any cell
    """
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(body_rows), cols=cols)
    table.autofit = True

    # Header
    hdr_cells = table.rows[0].cells
    for c_idx, txt in enumerate(header_row):
        cell = hdr_cells[c_idx]
        cell.text = ''
        p = cell.paragraphs[0]
        parse_inline_bold(p, txt)
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1F497D')
        set_cell_borders(cell, color='1F497D', size='6')

    # Body
    for r_idx, row in enumerate(body_rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx in range(cols):
            cell = cells[c_idx]
            cell.text = ''
            value = row[c_idx] if c_idx < len(row) else ''
            parse_inline_bold(cell.paragraphs[0], value)
            set_cell_shading(cell, 'DEEAF6')
            set_cell_borders(cell, color='BFBFBF', size='4')

    return table


def is_allcaps_heading(text):
    """
    Return True if the line looks like an ALL CAPS section heading.
    Matches lines like:
        WORK RESTRICTIONS:
        SUBJECTIVE COMPLAINTS/INTERIM HISTORY:
        RADIOGRAPHS:
        DATA:
        OBJECTIVE FINDINGS/PHYSICAL EXAMINATION:
        ASSESSMENT & PLAN:
    Does NOT match mixed-case lines, header fields with values on the same
    line (e.g. "DOB: 12/26/1993"), or short abbreviations (< 4 chars).
    """
    stripped = text.strip()
    if len(stripped) < 4:
        return False
    # All characters must be: uppercase A-Z, space, / & , ( ) . - :
    return bool(re.match(r'^[A-Z][A-Z\s/&,().:-]+:?\s*$', stripped))


# ---------------------------------------------------------------------------
# Main converter
# ---------------------------------------------------------------------------

HEADING_BLUE = RGBColor(0x1F, 0x49, 0x7D)

def convert_md_to_docx(md_path: Path, docx_path: Path):
    doc = Document()

    # --- Document-wide style tweaks ---
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    # Apply consistent Calibri + dark blue to all four heading levels
    heading_sizes = {
        'Heading 1': 16,
        'Heading 2': 13,
        'Heading 3': 12,
        'Heading 4': 11,
    }
    for h_style, size in heading_sizes.items():
        hs = doc.styles[h_style]
        hs.font.name = 'Calibri'
        hs.font.size = Pt(size)
        hs.font.color.rgb = HEADING_BLUE

    lines = md_path.read_text(encoding='utf-8').splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]

        # --- GFM pipe table (header + separator + body rows) ---
        if (is_table_row(line)
                and i + 1 < len(lines)
                and is_table_separator(lines[i + 1])):
            header = split_table_row(line)
            body = []
            j = i + 2
            while j < len(lines) and is_table_row(lines[j]) and not is_table_separator(lines[j]):
                body.append(split_table_row(lines[j]))
                j += 1
            add_md_table(doc, header, body)
            i = j
            continue

        # --- Heading 1: # Text ---
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)

        # --- Heading 2: ## Text ---
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)

        # --- Heading 3: ### Text ---
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)

        # --- Heading 4: #### Text ---
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)

        # --- Horizontal rule: --- ---
        elif line.strip() == '---':
            add_horizontal_rule(doc)

        # --- Bullet point: - Text ---
        elif re.match(r'^- ', line):
            content = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            parse_inline_bold(p, content)

        # --- Indented sub-bullet:   - Text (2+ leading spaces) ---
        elif re.match(r'^ {2,}- ', line):
            content = re.sub(r'^ +-\s*', '', line)
            p = doc.add_paragraph(style='List Bullet 2')
            parse_inline_bold(p, content)

        # --- Numbered list: 1. Text ---
        elif re.match(r'^\d+\.\s', line):
            content = re.sub(r'^\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            parse_inline_bold(p, content)

        # --- Blank line: small spacer paragraph ---
        elif line.strip() == '':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)

        # --- Normal paragraph or ALL CAPS heading ---
        else:
            stripped = line.strip()

            # ALL CAPS section heading (e.g. WC note sections like
            # "WORK RESTRICTIONS:", "SUBJECTIVE COMPLAINTS/INTERIM HISTORY:")
            if is_allcaps_heading(stripped):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(stripped)
                run.bold = True

            # Normal paragraph (may contain **bold** inline)
            else:
                p = doc.add_paragraph()
                parse_inline_bold(p, stripped)

        i += 1

    doc.save(str(docx_path))


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) != 2:
        print("Usage: python3 md_to_docx.py <path_to_soap_note.md>", file=sys.stderr)
        sys.exit(1)

    md_path = Path(sys.argv[1]).absolute()

    if not md_path.exists():
        print(f"ERROR: File not found: {md_path}", file=sys.stderr)
        sys.exit(1)

    if md_path.suffix != '.md':
        print(f"ERROR: Expected a .md file, got: {md_path.name}", file=sys.stderr)
        sys.exit(1)

    docx_path = md_path.with_suffix('.docx')

    try:
        convert_md_to_docx(md_path, docx_path)
    except Exception as e:
        print(f"ERROR: Conversion failed: {e}", file=sys.stderr)
        sys.exit(1)

    print(str(docx_path))


if __name__ == '__main__':
    main()
